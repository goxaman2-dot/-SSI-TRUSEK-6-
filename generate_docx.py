# -*- coding: utf-8 -*-
import os
import sys
import json
import math

# Try to install dependencies if not present (handled by server.ts, but safe to guard here)
try:
    import docx
    from docx.shared import Cm, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Error: python-docx is not installed.", file=sys.stderr)
    sys.exit(1)

try:
    import matplotlib
    matplotlib.use('Agg') # Non-interactive backend
    import matplotlib.pyplot as plt
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

def create_radar_chart(subfactors, output_img_path):
    if not HAS_MATPLOTLIB:
        return False
    
    try:
        # Define 6 main factors from the subfactors
        # Subfactors contains 12 elements. We can aggregate or show the 6 main dimensions:
        # T, R, U, S, E, K
        categories = ['T (Утилитарность)', 'R (Окупаемость)', 'U (Удержание)', 
                      'S (Виральность)', 'E (Драйв)', 'K (Капитал)']
        
        # We calculate the average score for each of the 6 factors
        # 12 subfactors are mapped as:
        # T: 0, 1
        # R: 2, 3
        # U: 4, 5
        # S: 6, 7
        # E: 8, 9
        # K: 10, 11
        scores = [0.0] * 6
        for sf in subfactors:
            code = sf.get('code', '')
            val = sf.get('value', 0)
            if code.startswith('T'):
                scores[0] += val / 2.0
            elif code.startswith('R'):
                scores[1] += val / 2.0
            elif code.startswith('U'):
                scores[2] += val / 2.0
            elif code.startswith('S'):
                scores[3] += val / 2.0
            elif code.startswith('E'):
                scores[4] += val / 2.0
            elif code.startswith('K'):
                scores[5] += val / 2.0
        
        # Number of variables
        N = len(categories)
        
        # What will be the angle of each axis in the plot? (we divide the circle / N)
        angles = [n / float(N) * 2 * math.pi for n in range(N)]
        angles += angles[:1]
        
        # Append first score to close the loop
        scores_closed = scores + scores[:1]
        
        # Initialise the spider plot
        fig, ax = plt.subplots(figsize=(5.5, 5.5), subplot_kw=dict(polar=True))
        
        # Draw one axe per variable + add labels
        plt.xticks(angles[:-1], categories, color='grey', size=8)
        
        # Draw ylabels
        ax.set_rlabel_position(0)
        plt.yticks([2, 4, 6, 8, 10], ["2", "4", "6", "8", "10"], color="grey", size=7)
        plt.ylim(0, 10)
        
        # Plot data
        ax.plot(angles, scores_closed, linewidth=2, linestyle='solid', color='#6366f1', label="Проект")
        
        # Fill area
        ax.fill(angles, scores_closed, '#818cf8', alpha=0.3)
        
        # Add a title
        plt.title('Лепестковая диаграмма устойчивости стартапа (TRUSEK-6)', size=11, color='#1e293b', y=1.1, weight='bold')
        
        # Save image
        plt.tight_layout()
        plt.savefig(output_img_path, dpi=150)
        plt.close()
        return True
    except Exception as e:
        print(f"Error drawing chart: {e}", file=sys.stderr)
        return False

def build_docx(data_json_path, output_docx_path):
    # Load input data
    with open(data_json_path, 'r', encoding='utf-8') as f:
        payload = json.load(f)
    
    data = payload.get('data', {})
    results = payload.get('results', {})
    
    startup_name = data.get('name', 'Безымянный стартап') or 'Безымянный стартап'
    author_name = data.get('author', 'Студент СКФУ') or 'Студент СКФУ'
    mentor_name = data.get('expert', 'Наставник') or 'Наставник'
    final_ssi = float(results.get('finalSsi', 0.0) or 0.0)
    interpretation = results.get('interpretation', 'Не определено') or 'Не определено'
    
    # Reconstruct the 12 subfactors list exactly as in the UI
    def norm_val(val, min_val, max_val, invert=False):
        if val is None or val == "":
            return 0.0
        try:
            val = float(val)
        except (ValueError, TypeError):
            return 0.0
        if max_val == min_val:
            return 0.0
        v = ((val - min_val) / (max_val - min_val)) * 10.0
        if invert:
            v = 10.0 - v
        return max(0.0, min(10.0, v))

    def get_raw_val(code):
        try:
            if code == 'T1':
                return f"{data.get('t1', 0)} мес."
            elif code == 'T2':
                return f"{data.get('t2', 0)} мес."
            elif code == 'U1':
                u1 = float(data.get('u1', 0) or 0)
                if u1 >= 1000:
                    return f"{u1 / 1000:.0f} тыс.р."
                return f"{u1:.0f} руб"
            elif code == 'U2':
                return f"{data.get('u2', 0)} раз/г"
            elif code == 'R1':
                return f"{data.get('r1', 0)} р/г"
            elif code == 'R2':
                return f"{float(data.get('r2', 1.0) or 1.0):.1f}x"
            elif code == 'S1':
                return f"{data.get('s1', 0)}%"
            elif code == 'S2':
                return f"{data.get('s2', 0)}"
            elif code == 'E1':
                return f"{data.get('e1', 0)} мин."
            elif code == 'E2':
                return f"{float(data.get('e2', 1.0) or 1.0):.2f}x"
            elif code == 'K1':
                k1 = float(data.get('k1', 0) or 0)
                if k1 >= 1000:
                    return f"{k1 / 1000:.1f} млн"
                return f"{k1:.0f} тыс.р."
            elif code == 'K2':
                return f"{data.get('k2', 0)}%"
        except Exception:
            pass
        return str(data.get(code.lower(), ''))

    subfactors = [
        {
            'code': 'T1',
            'name': 'До EBITDA+',
            'value': norm_val(data.get('t1'), 1, 36, invert=True),
            'raw_val': get_raw_val('T1')
        },
        {
            'code': 'T2',
            'name': 'Окупаемость',
            'value': norm_val(data.get('t2'), 1, 60, invert=True),
            'raw_val': get_raw_val('T2')
        },
        {
            'code': 'U1',
            'name': 'Ущерб отказа',
            'value': norm_val(data.get('u1'), 0, 300000),
            'raw_val': get_raw_val('U1')
        },
        {
            'code': 'U2',
            'name': 'Потребность',
            'value': norm_val(data.get('u2'), 1, 52),
            'raw_val': get_raw_val('U2')
        },
        {
            'code': 'R1',
            'name': 'Частота покуп',
            'value': norm_val(data.get('r1'), 0, 52),
            'raw_val': get_raw_val('R1')
        },
        {
            'code': 'R2',
            'name': 'LTV/CAC',
            'value': norm_val(data.get('r2'), 0, 10),
            'raw_val': get_raw_val('R2')
        },
        {
            'code': 'S1',
            'name': 'Рекомендац.',
            'value': norm_val(data.get('s1'), 0, 100),
            'raw_val': get_raw_val('S1')
        },
        {
            'code': 'S2',
            'name': 'Лояльность NPS',
            'value': norm_val(data.get('s2'), -100, 100),
            'raw_val': get_raw_val('S2')
        },
        {
            'code': 'E1',
            'name': 'Время сессии',
            'value': norm_val(data.get('e1'), 0, 60),
            'raw_val': get_raw_val('E1')
        },
        {
            'code': 'E2',
            'name': 'Качество/Цена',
            'value': norm_val(data.get('e2'), 0, 2),
            'raw_val': get_raw_val('E2')
        },
        {
            'code': 'K1',
            'name': 'Старт CAPEX',
            'value': norm_val(data.get('k1'), 0, 5000, invert=True),
            'raw_val': get_raw_val('K1')
        },
        {
            'code': 'K2',
            'name': 'Маржа OPEX',
            'value': norm_val(data.get('k2'), 0, 80),
            'raw_val': get_raw_val('K2')
        }
    ]
    
    doc = docx.Document()
    
    # Define custom styles
    # Set default margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        
    # Helper to apply standard Russian academic paragraph settings:
    # 1.5 Line spacing, 1.25 cm first line indent, justified alignment, default font Inter/Times
    def add_academic_paragraph(text="", bold=False, italic=False, align_justify=True, first_indent=True):
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.line_spacing = 1.5
        if first_indent:
            p_format.first_line_indent = Cm(1.25)
        else:
            p_format.first_line_indent = Cm(0)
            
        if align_justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        p_format.space_after = Pt(6)
        p_format.space_before = Pt(0)
        
        if text:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = bold
            run.italic = italic
        return p

    def add_heading_1(text):
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.line_spacing = 1.5
        p_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_format.space_before = Pt(12)
        p_format.space_after = Pt(6)
        p_format.keep_with_next = True
        
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.bold = True
        return p

    # --- TITLE ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(50)
    p_title.paragraph_format.space_after = Pt(24)
    r_title = p_title.add_run("НАУЧНО-ПРАКТИЧЕСКИЙ ОТЧЕТ\nОЦЕНКИ ЖИЗНЕСПОСОБНОСТИ СТАРТАП-ПРОЕКТА")
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(16)
    r_title.bold = True

    # Subtitle with project name
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(36)
    r_sub = p_sub.add_run(f"«Расчет индекса самодостаточности SSI (Self-Sufficiency Index) для стартапа «{startup_name}» по методике TRUSEK-6 + MEI»")
    r_sub.font.name = 'Times New Roman'
    r_sub.font.size = Pt(14)
    r_sub.italic = True
    r_sub.bold = True

    # Meta-information (Right aligned)
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_meta.paragraph_format.space_after = Pt(48)
    r_meta = p_meta.add_run(
        f"Выполнил студент-разработчик: {author_name}\n"
        f"Научный наставник / эксперт: {mentor_name}\n"
        f"Итоговый индекс устойчивости SSI: {final_ssi:.2f} / 10.0\n"
        f"Интерпретация: {interpretation.split('—')[0]}\n"
    )
    r_meta.font.name = 'Times New Roman'
    r_meta.font.size = Pt(12)
    r_meta.bold = True

    doc.add_page_break()

    # --- SECTION 1: INTRODUCTION ---
    add_heading_1("1. Введение и научная методология исследования")
    add_academic_paragraph(
        "В современных условиях академического предпринимательства и развития молодежных инноваций "
        "критически важной задачей является ранняя диагностика жизнеспособности коммерческих проектов. "
        "Настоящий научно-практический отчет составлен на основе применения авторской методологии "
        "расчета комплексного индекса устойчивости стартап-проектов SSI (Self-Sufficiency Index), "
        "включающего двенадцатифакторную модель TRUSEK-6, адаптированную для оценки рыночного "
        "модификатора MEI (Market Efficiency Index). Данная методика позволяет всесторонне оценить "
        "внутренний технологический, экономический, маркетинговый и управленческий потенциал стартапа."
    )
    add_academic_paragraph(
        f"В качестве объекта исследования выступает стартап-проект «{startup_name}», разрабатываемый "
        f"студентом {author_name} под научно-методическим руководством наставника {mentor_name}. "
        "Расчет произведен в интерактивной среде интеллектуального калькулятора, осуществляющего "
        "автоматизированную проверку согласованности рыночных и финансовых гипотез."
    )

    # --- SECTION 2: CALCULATION RESULTS ---
    add_heading_1("2. Анализ итоговых показателей и факторов устойчивости")
    add_academic_paragraph(
        f"По результатам проведенного расчета, итоговый индекс самодостаточности SSI проекта "
        f"«{startup_name}» составил {final_ssi:.2f} из 10.0 возможных баллов. Согласно шкале "
        f"интерпретации авторов методики, данный уровень характеризуется как: «{interpretation}»."
    )
    add_academic_paragraph(
        "Каждый из шести ключевых факторов (Утилитарность, Окупаемость, Удержание, Виральность, "
        "Энергия/Драйв и Капитал) имеет определенный весовой коэффициент, установленный экспертным путем "
        "или настроенный индивидуально в соответствии со спецификой технологического сектора проекта."
    )

    # TABLE 1 Caption (Above and Left aligned according to request: "к таблицам слева")
    p_cap1 = doc.add_paragraph()
    p_cap1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_cap1.paragraph_format.first_line_indent = Cm(0)
    p_cap1.paragraph_format.space_before = Pt(12)
    p_cap1.paragraph_format.space_after = Pt(4)
    r_cap1 = p_cap1.add_run(f"Таблица 1 – Результаты расчета факторов индекса устойчивости SSI стартапа «{startup_name}»")
    r_cap1.font.name = 'Times New Roman'
    r_cap1.font.size = Pt(12)
    r_cap1.bold = True

    # TABLE 1
    # We will aggregate subfactors to 6 main factors
    factor_names = {
        'T': 'T · Утилитарность (Utility & Problem)',
        'R': 'R · Окупаемость (Payback & Unit)',
        'U': 'U · Удержание клиентов (Retention)',
        'S': 'S · Виральность и Охват (Virality)',
        'E': 'E · Энергия и Драйв команды (Energy)',
        'K': 'K · Капитал и Финансы (Capital)'
    }
    
    # Calculate factor aggregates with exact default weights matching the SSI formula
    factor_weights_defaults = {
        'U': 0.15,
        'E': 0.20,
        'R': 0.15,
        'K': 0.15,
        'T': 0.20,
        'S': 0.15
    }
    factors_data = {k: {'val': 0.0, 'weight': factor_weights_defaults.get(k, 0.1667)} for k in factor_names.keys()}
    # Load custom weights if available in metadata
    weights_map = data.get('weights', {}) # {t: 20, r: 20...}
    if weights_map:
        for k in factor_names.keys():
            lower_k = k.lower()
            if lower_k in weights_map:
                factors_data[k]['weight'] = weights_map[lower_k] / 100.0

    for sf in subfactors:
        code = sf.get('code', '')
        val = sf.get('value', 0)
        prefix = code[0] if code else ''
        if prefix in factors_data:
            factors_data[prefix]['val'] += val / 2.0 # Since 2 subfactors per factor

    table1 = doc.add_table(rows=1, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.LEFT # "таблицы от левого края страницы"
    
    # Header row
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Обозначение'
    hdr_cells[1].text = 'Наименование фактора модели TRUSEK-6'
    hdr_cells[2].text = 'Вес (%)'
    hdr_cells[3].text = 'Оценка (0-10)'
    
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
                
    # Populate rows
    for key, f_info in factors_data.items():
        row_cells = table1.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = factor_names[key]
        row_cells[2].text = f"{f_info['weight']*100:.1f}%"
        row_cells[3].text = f"{f_info['val']:.2f}"
        
        # Style text in cells
        for idx, cell in enumerate(row_cells):
            p = cell.paragraphs[0]
            if idx in [0, 2, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    add_academic_paragraph(
        "Представленная таблица демонстрирует дифференцированный вклад каждого измерения в итоговую "
        "стабильность стартап-модели. Данные служат основой для выявления диспропорций в бизнес-плане.",
        first_indent=True
    )

    # --- SECTION 3: GRAPHICS ---
    add_heading_1("3. Визуализация устойчивости (Лилия SSI)")
    add_academic_paragraph(
        "Для наглядного представления сильных сторон проекта и уязвимых зон в научно-практическом отчете "
        "используется геометрическая интерпретация лепестков устойчивости (Лилия SSI), построенная в виде "
        "лепестковой (радарной) диаграммы. Чем шире площадь лепестка, тем устойчивее соответствующее "
        "направление бизнеса."
    )

    # Draw radar chart
    img_path = "temp_radar_chart.png"
    chart_success = create_radar_chart(subfactors, img_path)
    
    if chart_success and os.path.exists(img_path):
        # Insert image (Centered as requested: "рисунки по центру")
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(12)
        p_img.paragraph_format.space_after = Pt(6)
        p_img.paragraph_format.first_line_indent = Cm(0)
        
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(4.5))
        
        # Caption BELOW the picture as requested: "названия к рисункам снизу"
        p_cap2 = doc.add_paragraph()
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap2.paragraph_format.first_line_indent = Cm(0)
        p_cap2.paragraph_format.space_after = Pt(12)
        r_cap2 = p_cap2.add_run(f"Рисунок 1 – Лепестковая диаграмма факторов устойчивости (Лилия SSI) стартапа «{startup_name}»")
        r_cap2.font.name = 'Times New Roman'
        r_cap2.font.size = Pt(11)
        r_cap2.italic = True
        r_cap2.bold = True
        
        # Delete temporary image
        try:
            os.remove(img_path)
        except OSError:
            pass
    else:
        add_academic_paragraph(
            "[Внимание: Графическое изображение временно недоступно. Для полноценной отрисовки "
            "лепестковой диаграммы требуется установка графического пакета matplotlib в серверной среде.]",
            italic=True
        )

    # --- SECTION 4: SUBFACTORS ---
    add_heading_1("4. Детализированная оценка подфакторов и метрик")
    add_academic_paragraph(
        "Каждый макрофактор модели TRUSEK-6 раскрывается через два специализированных подфактора, "
        "оценивающих конкретные показатели операционной деятельности, маркетинга и структуры команды. "
        "Ниже приведена детальная спецификация всех двенадцати показателей проекта."
    )

    # TABLE 2 Caption
    p_cap3 = doc.add_paragraph()
    p_cap3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_cap3.paragraph_format.first_line_indent = Cm(0)
    p_cap3.paragraph_format.space_before = Pt(12)
    p_cap3.paragraph_format.space_after = Pt(4)
    r_cap3 = p_cap3.add_run(f"Таблица 2 – Детализированные оценки подфакторов стартапа «{startup_name}»")
    r_cap3.font.name = 'Times New Roman'
    r_cap3.font.size = Pt(12)
    r_cap3.bold = True

    # TABLE 2
    table2 = doc.add_table(rows=1, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.LEFT # Left aligned table
    
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Код'
    hdr2[1].text = 'Наименование подфактора'
    hdr2[2].text = 'Оценка'
    hdr2[3].text = 'Текущий статус / Значение'
    
    for cell in hdr2:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True

    # Format values for readability
    for sf in subfactors:
        row_cells = table2.add_row().cells
        row_cells[0].text = sf.get('code', '')
        row_cells[1].text = sf.get('name', '')
        row_cells[2].text = f"{sf.get('value', 0):.1f}"
        
        desc_text = sf.get('raw_val', '')
        row_cells[3].text = desc_text
        
        for idx, cell in enumerate(row_cells):
            p = cell.paragraphs[0]
            if idx in [0, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    # --- SECTION 5: RECOMMENDATIONS ---
    add_heading_1("5. Заключение и адресные рекомендации ментора")
    add_academic_paragraph(
        f"Комплексный анализ показателей стартапа «{startup_name}» свидетельствует о наличии "
        "как сильных конкурентных преимуществ, так и зон высокого риска, требующих оперативного "
        "управленческого вмешательства."
    )
    
    # Generate custom advice for weak subfactors
    weak_factors = [sf for sf in subfactors if sf.get('value', 10) < 5.0]
    if weak_factors:
        add_academic_paragraph(
            "В ходе анализа были зафиксированы следующие дефицитные зоны устойчивости (оценка ниже 5.0 баллов):",
            bold=True
        )
        for wf in weak_factors:
            p_bullet = doc.add_paragraph()
            p_bullet.paragraph_format.line_spacing = 1.5
            p_bullet.paragraph_format.first_line_indent = Cm(1.25)
            p_bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_bullet.paragraph_format.space_after = Pt(4)
            
            run_wf = p_bullet.add_run(f"• Фактор {wf.get('code', '')} ({wf.get('name', '')}) — оценка {wf.get('value', 0):.1f} баллов. ")
            run_wf.font.name = 'Times New Roman'
            run_wf.font.size = Pt(13)
            run_wf.bold = True
            
            # Context recommendations
            rec_text = "Рекомендуется пересмотреть ценностное предложение и оптимизировать операционные расходы."
            code = wf.get('code', '')
            if code == 'T1':
                rec_text = "Сфокусируйтесь на более острой боли пользователей. Проведите дополнительные глубинные интервью (CustDev) для сужения сегмента."
            elif code == 'T2':
                rec_text = "Высокая сила альтернативных решений. Найдите уникальное ценностное отличие (УТП), которое трудно скопировать конкурентам."
            elif code == 'R1':
                rec_text = "Маржинальность под угрозой. Снижайте стоимость привлечения клиента (CAC) или увеличивайте его пожизненную ценность (LTV) за счет допродаж."
            elif code == 'R2':
                rec_text = "Слишком долгий срок окупаемости. Ускорьте первый цикл сделки, внедрите предоплату или упростите продукт до MVP."
            elif code == 'U1' or code == 'U2':
                rec_text = "Низкое удержание. Внедрите систему триггерных рассылок, улучшите онбординг пользователей, собирайте обратную связь при оттоке клиентов."
            elif code == 'S1':
                rec_text = "Слабый вирусный рост. Добавьте реферальные программы («приведи друга»), геймификацию или бесшовный экспорт результатов работы в соцсети."
            elif code == 'S2':
                rec_text = "Малый объем доступного рынка. Рассмотрите смежные сегменты клиентов, диверсифицируйте продуктовую линейку или масштабируйте географию."
            elif code == 'E1' or code == 'E2':
                rec_text = "Дефицит компетенций или энергии. Привлеките недостающих специалистов в команду за долю в проекте, проведите тимбилдинг и распределите зоны ответственности."
            elif code == 'K1' or code == 'K2':
                rec_text = "Острый дефицит финансовой автономии. Сократите постоянные издержки (Burn Rate), рассмотрите привлечение грантов (например, Студенческий стартап Фонда содействия инновациям) или предпосевных инвестиций."
                
            run_rec = p_bullet.add_run(rec_text)
            run_rec.font.name = 'Times New Roman'
            run_rec.font.size = Pt(13)
            
    else:
        add_academic_paragraph(
            "Поздравляем! Все исследуемые подфакторы стартапа находятся в зоне высокой и средней устойчивости "
            "(выше 5.0 баллов). Проекту рекомендуется удерживать текущие операционные темпы и готовиться к "
            "масштабированию коммерческой деятельности."
        )

    # Mentor Sign-off
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign.paragraph_format.space_before = Pt(36)
    r_sign = p_sign.add_run(
        f"Отчет сформирован автоматически в системе «Индекс SSI»\n"
        f"Подпись наставника / эксперта: ____________________ / {mentor_name} /"
    )
    r_sign.font.name = 'Times New Roman'
    r_sign.font.size = Pt(12)
    r_sign.italic = True

    # Save document
    doc.save(output_docx_path)
    print("Document successfully generated!")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python3 generate_docx.py <input_json_path> <output_docx_path>")
        sys.exit(1)
        
    build_docx(sys.argv[1], sys.argv[2])
